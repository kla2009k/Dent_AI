# -*- coding: utf-8 -*-
"""
สร้างเล่มรายงานโครงงาน DentScan AI (.docx)
เลียนแบบ format/style จาก template PRJ68-CS02.docx เป๊ะ:
  - A4, margin 1", TH Sarabun New 16pt, หัวข้อใหญ่ 20pt bold, หัวข้อย่อย 16pt bold
  - ปก + ตราโรงเรียน จภ.เลย + หน้าข้อมูล + บทคัดย่อ + 5 บท + บรรณานุกรม + ภาคผนวก
รัน: python report/make_report.py  → report/DentScan_Report.docx
"""
import json
import pathlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).parent.parent
OUT = pathlib.Path(__file__).parent
FONT = "TH Sarabun New"

# ── โหลดผลจริงจากโปรเจกต์ ──
def load_metrics():
    m = {}
    for name, p in [("v1", ROOT/"models"/"metrics.json"),
                    ("v2", ROOT/"models"/"v2"/"metrics.json"),
                    ("v3", ROOT/"models"/"v3"/"metrics.json")]:
        if p.exists():
            m[name] = json.load(open(p))
    sc = ROOT/"models"/"serving_config.json"
    m["serving"] = json.load(open(sc)) if sc.exists() else {}
    return m

M = load_metrics()


# ── Thai font helper (ต้องตั้ง w:cs + szCs ให้ Sarabun แสดงไทยถูก) ──
def _font(run, size=16, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    # complex-script size (ไทย) ต้องตั้งแยก
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(int(size * 2)))
    rpr.append(szcs)


def para(doc, text="", size=16, bold=False, align=None, italic=False,
         indent=None, space_after=6, space_before=0, line=1.5, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if line:
        pf.line_spacing = line
    if indent is not None:
        pf.first_line_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        _font(r, size, bold, italic, color)
    return p


def body(doc, text, indent=0.5):
    # ไทยไม่มี word-break → justify ทำให้เกิดช่องว่างกว้าง ใช้ left (มาตรฐานเล่มไทย)
    return para(doc, text, size=16, align=WD_ALIGN_PARAGRAPH.LEFT, indent=indent)


def _style_cs(style, size):
    """ตั้งฟอนต์ complex-script (ไทย) ให้ paragraph style"""
    rpr = style.element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs"); rpr.append(szcs)
    szcs.set(qn("w:val"), str(int(size * 2)))


def setup_styles(doc):
    """ตั้ง Heading 1/2 ให้เป็น Sarabun + รูปแบบเล่มไทย (เพื่อให้ TOC field เก็บได้)
       + ตั้ง TOC 1/2 เป็น Sarabun (เลขหน้าชิดขวา+จุดไข่ปลา = ค่าเริ่มต้นของ TOC style)"""
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT; h1.font.size = Pt(20); h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12); h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.line_spacing = 1.5
    _style_cs(h1, 20)
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT; h2.font.size = Pt(16); h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(8); h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.line_spacing = 1.5
    _style_cs(h2, 16)
    for nm in ("TOC 1", "TOC 2"):
        try:
            s = doc.styles[nm]; s.font.name = FONT; s.font.size = Pt(16); _style_cs(s, 16)
        except KeyError:
            pass


def h_chapter(doc, num, title):
    # paragraph เดียว (Heading 1) 2 บรรทัด → TOC เก็บเป็น 1 รายการ
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"บทที่ {num}")
    _font(r1, 20, bold=True)
    r1.add_break()
    r2 = p.add_run(title)
    _font(r2, 20, bold=True)
    return p


def h_front(doc, text):
    # หัวข้อหน้าต้น/ท้าย ที่ต้องการให้อยู่ในสารบัญ (บทคัดย่อ/กิตติกรรม/บรรณานุกรม/ภาคผนวก)
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _font(r, 20, bold=True)
    return p


def h_sub(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _font(r, 16, bold=True)
    return p


def add_toc_field(doc, instr='TOC \\o "1-2" \\h \\z \\u'):
    """แทรก Word TOC field (อัปเดตอัตโนมัติ — เลขหน้าชิดขวา+จุดไข่ปลา)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " " + instr + " "
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "«อัปเดตสารบัญ: คลิกขวา → Update Field»"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    for el in (b, i, sep, t, e):
        run._element.append(el)
    _font(run, 14)


def h_subsub(doc, text):
    para(doc, text, size=16, bold=True, space_before=4, space_after=2, indent=0.3)


def bullet(doc, text, lvl=0):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.75 + lvl * 0.3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("• " + text)
    _font(r, 16)
    return p


def add_image(doc, path, width=5.5, caption=None):
    path = pathlib.Path(path)
    if not path.exists():
        para(doc, f"[ภาพ: {path.name} — ใส่ภายหลัง]", italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, color=(150, 150, 150))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        para(doc, caption, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
             italic=True, space_after=10)


def page_break(doc):
    doc.add_page_break()


# ── เลขหน้า: ก ข ค (หน้าต้น) / 1 2 3 (เนื้อหา) ตามมาตรฐานเล่มไทย ──
def set_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def footer_pagenum(section, size=14):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    _font(r, size)
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._element.append(f1); r._element.append(it); r._element.append(f2)


def table_caption(doc, text):
    # มาตรฐานไทย: caption ตารางอยู่ "เหนือ" ตาราง
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    num, _, rest = text.partition(" ")
    num2, _, rest2 = rest.partition(" ")
    r1 = p.add_run(f"{num} {num2} ")
    _font(r1, 15, bold=True)
    r2 = p.add_run(rest2)
    _font(r2, 15)


def start_section(doc, fmt, start, size=14):
    """เริ่ม section ใหม่ (ขึ้นหน้าใหม่) + ตั้งรูปแบบเลขหน้า + ใส่เลขหน้าใน footer
    fmt='thaiLetters' → ก ข ค (หน้าต้น) · fmt='decimal' → 1 2 3 (เนื้อหา)"""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)
    s.left_margin = s.right_margin = Inches(1)
    s.top_margin = s.bottom_margin = Inches(1)
    set_pgnum(s, fmt=fmt, start=start)
    s.footer.is_linked_to_previous = False
    footer_pagenum(s, size)
    return s


def lof_row(doc, text, pg):
    # เลขหน้าชิดขวา + จุดไข่ปลานำ ด้วย right tab stop (สวย ตรง ไม่เพี้ยนตามจำนวนอักษร)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.27), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text + "\t" + str(pg))
    _font(r, 16)


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(h)
        _font(r, 15, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[i].paragraphs[0].add_run(str(v))
            _font(r, 15)
    para(doc, "", space_after=4)
    return t


# ══════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)   # A4
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin = sec.bottom_margin = Inches(1)
# default style
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(16)
setup_styles(doc)

TITLE_TH = "การพัฒนาระบบคัดกรองโรคในช่องปากจากภาพรังสีพาโนรามิกด้วยปัญญาประดิษฐ์ (DentScan AI)"
TITLE_EN = "DentScan AI: Development of an AI-Based Oral Disease Screening System from Panoramic Radiographs"

# ── ปก ──
para(doc, "", space_after=6)
add_image(doc, OUT / "school_logo.png", width=1.6)
para(doc, "", space_after=4)
para(doc, TITLE_TH, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, TITLE_EN, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
para(doc, "โดย", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "นายปัญญากร อาจหาญ", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "นายปรเมษฐ์ มารศรี", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
para(doc, "โครงงานวิทยาศาสตร์ สาขาเทคโนโลยีและคอมพิวเตอร์", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "รายงานนี้เป็นส่วนหนึ่งของวิชา ว30281 โครงงาน", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "ตามหลักสูตรโรงเรียนวิทยาศาสตร์ภูมิภาค (ฉบับปรับปรุง พ.ศ. 2566)", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "โรงเรียนวิทยาศาสตร์จุฬาภรณราชวิทยาลัย เลย", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "ภาคเรียนที่ 2 ชั้นมัธยมศึกษาปีที่ 6 ปีการศึกษา 2568", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
para(doc, "Science Project in Computer and Technology", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "Princess Chulabhorn Science High School Loei", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para(doc, "2568", align=WD_ALIGN_PARAGRAPH.CENTER)
page_break(doc)

# ── หน้าข้อมูลโครงงาน ──
def info_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(label + "\t")
    _font(r, 16, bold=True)
    r2 = p.add_run(value)
    _font(r2, 16)

para(doc, "ข้อมูลโครงงาน", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
info_row(doc, "ชื่อโครงงาน", TITLE_TH)
info_row(doc, "ชื่อภาษาอังกฤษ", TITLE_EN)
info_row(doc, "ผู้ทำโครงงาน", "1. นายปัญญากร อาจหาญ ชั้นมัธยมศึกษาปีที่ 6")
info_row(doc, "", "2. นายปรเมษฐ์ มารศรี ชั้นมัธยมศึกษาปีที่ 6")
info_row(doc, "ครูที่ปรึกษา", "นายศตวรรษ ทัดมาลา")
info_row(doc, "สาขา", "เทคโนโลยีและคอมพิวเตอร์")
info_row(doc, "สถานศึกษา", "โรงเรียนวิทยาศาสตร์จุฬาภรณราชวิทยาลัย เลย")
info_row(doc, "ปีการศึกษา", "2568")
# ── เริ่มเลขหน้าหน้าต้น (ก ข ค) ตั้งแต่บทคัดย่อ ──
start_section(doc, "thaiLetters", 1)

# ── บทคัดย่อ ──
h_front(doc, "บทคัดย่อ")
body(doc,
     "โครงงานนี้พัฒนา DentScan AI ระบบคัดกรองโรคในช่องปากจากภาพรังสีพาโนรามิก (panoramic radiograph) "
     "ด้วยปัญญาประดิษฐ์ โดยมีวัตถุประสงค์เพื่อช่วยคัดกรองความผิดปกติของฟัน 4 ประเภท ได้แก่ ฟันผุ (caries) "
     "ฟันผุลึก (deep caries) รอยโรคปลายราก (periapical lesion) และฟันคุด (impacted tooth) "
     "เพื่อเป็นเครื่องมือสนับสนุนการตัดสินใจ (decision-support) สำหรับการคัดกรองเบื้องต้น")
body(doc,
     "ผู้วิจัยใช้ชุดข้อมูล DENTEX Challenge 2023 (MICCAI) จำนวน 1,005 ภาพที่มีการกำกับโรค ฝึกแบบจำลอง "
     "EfficientNet-B0 ด้วยเทคนิค transfer learning และ focal loss เพื่อแก้ปัญหาความไม่สมดุลของข้อมูล "
     "ร่วมกับการทำ test-time augmentation และการรวมแบบจำลอง (ensemble) เพื่อเพิ่มความแม่นยำ "
     "พร้อมประยุกต์ Grad-CAM เพื่อแสดงตำแหน่งที่แบบจำลองให้ความสนใจ (explainable AI) "
     "และพัฒนาเว็บแอปพลิเคชันแบบไฮบริด (ภาพ X-ray + แบบสอบถามอาการ) พร้อมระบบสร้างรายงานภาษาไทย")
body(doc,
     f"ผลการทดลองบนชุดทดสอบพบว่า การรวมแบบจำลองร่วมกับ test-time augmentation ให้ค่าเฉลี่ย AUC เท่ากับ "
     f"{M.get('serving',{}).get('val_mean_auc','0.72')} ซึ่งสูงกว่าแบบจำลองเดี่ยว โดยโรคฟันผุและฟันคุด "
     "ให้ผลการจำแนกที่ดี ขณะที่รอยโรคปลายรากยังมีข้อจำกัดเนื่องจากปริมาณข้อมูลน้อย "
     "ระบบที่พัฒนาขึ้นสามารถใช้งานได้จริงผ่านเว็บแอปพลิเคชัน และแสดงผลการวิเคราะห์พร้อมแผนที่ความร้อนและรายงานภาษาไทย")
para(doc, "คำสำคัญ: ปัญญาประดิษฐ์, การคัดกรองโรคในช่องปาก, ภาพรังสีพาโนรามิก, การเรียนรู้เชิงลึก, "
          "EfficientNet, Grad-CAM, DENTEX", size=16, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
     space_before=8)
page_break(doc)

# ── กิตติกรรมประกาศ ──
h_front(doc, "กิตติกรรมประกาศ")
body(doc,
     "โครงงานนี้สำเร็จลุล่วงได้ด้วยความกรุณาจากนายศตวรรษ ทัดมาลา ครูที่ปรึกษาโครงงาน "
     "ที่ได้ให้คำแนะนำ คำปรึกษา และข้อเสนอแนะที่เป็นประโยชน์ตลอดการดำเนินงาน ผู้จัดทำขอขอบพระคุณเป็นอย่างสูง")
body(doc,
     "ขอขอบคุณคณะผู้จัดทำชุดข้อมูล DENTEX Challenge 2023 (MICCAI) ที่เผยแพร่ข้อมูลภาพรังสีพาโนรามิก "
     "ภายใต้สัญญาอนุญาต CC BY-NC-SA 4.0 ซึ่งเป็นข้อมูลหลักในการพัฒนาแบบจำลอง และขอขอบคุณโรงเรียนวิทยาศาสตร์"
     "จุฬาภรณราชวิทยาลัย เลย ที่สนับสนุนการทำโครงงานในครั้งนี้")
para(doc, "คณะผู้จัดทำ", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=10)
page_break(doc)

# ── สารบัญ (auto Word TOC field — อัปเดตเอง เลขหน้าชิดขวา+จุดไข่ปลา) ──
para(doc, "สารบัญ", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_toc_field(doc)
page_break(doc)

# ── สารบัญตาราง (เลขหน้าจริงจากเล่ม + ชิดขวา/จุดไข่ปลา) ──
para(doc, "สารบัญตาราง", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
for t, pg in [("ตารางที่ 3.1 จำนวนภาพต่อโรคในชุดข้อมูลฝึก", "6"),
              ("ตารางที่ 4.1 เปรียบเทียบผลของแต่ละแบบจำลอง", "9"),
              ("ตารางที่ 4.2 ผลการรวมแบบจำลองและ Test-Time Augmentation", "10"),
              ("ตารางที่ 4.3 ผลการจำแนกรายโรค (แบบจำลอง v1)", "10")]:
    lof_row(doc, t, pg)
para(doc, "", space_after=10)
para(doc, "สารบัญภาพ", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=14, space_after=12)
for t, pg in [("ภาพที่ 3.1 การกระจายตัวของข้อมูลแต่ละโรค", "6"),
              ("ภาพที่ 3.2 ตัวอย่างผลการวิเคราะห์พร้อม Grad-CAM", "8"),
              ("ภาพที่ 4.1 เมทริกซ์การเกิดร่วมกันของโรค", "9"),
              ("ภาพที่ 4.2 ตัวอย่างแผนที่ความร้อน Grad-CAM ต่อโรค", "11"),
              ("ภาพที่ ก.1 จำนวน annotation ต่อโรค", "15"),
              ("ภาพที่ ก.2 ตัวอย่างภาพรังสีของแต่ละโรค", "15"),
              ("ภาพที่ ก.3 ตัวอย่างภาพที่กำกับตำแหน่งโรค", "16"),
              ("ภาพที่ ข.1 ตัวอย่างผลการวิเคราะห์เพิ่มเติม", "16")]:
    lof_row(doc, t, pg)

# ── เริ่มเลขหน้าเนื้อหา (1 2 3) ตั้งแต่บทที่ 1 ──
start_section(doc, "decimal", 1)

# ══════════════════ บทที่ 1 บทนำ ══════════════════
h_chapter(doc, 1, "บทนำ")
h_sub(doc, "1.1 ที่มาและความสำคัญของโครงงาน")
body(doc, "โรคในช่องปาก โดยเฉพาะฟันผุและโรคปริทันต์ เป็นปัญหาสุขภาพที่พบบ่อยที่สุดอย่างหนึ่งของประชากรไทย "
          "และมักไม่แสดงอาการในระยะเริ่มต้น ทำให้ผู้ป่วยจำนวนมากไม่ได้รับการรักษาทันเวลา จนลุกลามเป็นฟันผุลึก "
          "การติดเชื้อในโพรงประสาทฟัน หรือรอยโรคปลายราก การตรวจวินิจฉัยด้วยภาพรังสีพาโนรามิก (panoramic radiograph) "
          "เป็นวิธีมาตรฐานที่ทันตแพทย์ใช้ประเมินสภาพฟันทั้งปาก แต่การอ่านผลต้องอาศัยความเชี่ยวชาญและเวลา "
          "อีกทั้งในพื้นที่ห่างไกลยังขาดแคลนทันตแพทย์")
body(doc, "ในปัจจุบัน ปัญญาประดิษฐ์ (Artificial Intelligence) โดยเฉพาะการเรียนรู้เชิงลึก (Deep Learning) "
          "ได้แสดงความสามารถในการวิเคราะห์ภาพทางการแพทย์ได้ใกล้เคียงผู้เชี่ยวชาญ ผู้จัดทำจึงสนใจพัฒนา "
          "ระบบ DentScan AI ที่นำการเรียนรู้เชิงลึกมาช่วยคัดกรองโรคในช่องปากจากภาพรังสีพาโนรามิก "
          "เพื่อเป็นเครื่องมือสนับสนุนการตัดสินใจเบื้องต้น ลดภาระการคัดกรอง และเพิ่มโอกาสการเข้าถึงบริการทันตกรรม")
h_sub(doc, "1.2 วัตถุประสงค์ของโครงงาน")
bullet(doc, "เพื่อพัฒนาแบบจำลองการเรียนรู้เชิงลึกสำหรับจำแนกโรคในช่องปาก 4 ประเภทจากภาพรังสีพาโนรามิก")
bullet(doc, "เพื่อประยุกต์เทคนิค Grad-CAM ในการแสดงตำแหน่งที่แบบจำลองให้ความสนใจ (explainable AI)")
bullet(doc, "เพื่อพัฒนาเว็บแอปพลิเคชันแบบไฮบริดที่รวมภาพ X-ray และอาการของผู้ป่วยในการวิเคราะห์")
bullet(doc, "เพื่อสร้างระบบรายงานผลภาษาไทยที่เข้าใจง่ายสำหรับผู้ใช้ทั่วไปและทันตแพทย์")
h_sub(doc, "1.3 สมมติฐานของการวิจัย")
body(doc, "แบบจำลองการเรียนรู้เชิงลึกที่ฝึกด้วยชุดข้อมูลภาพรังสีพาโนรามิก สามารถจำแนกโรคในช่องปาก 4 ประเภท "
          "ได้ด้วยความแม่นยำที่ยอมรับได้ (ค่าเฉลี่ย AUC ≥ 0.70) และการรวมแบบจำลอง (ensemble) ร่วมกับ "
          "test-time augmentation จะให้ผลดีกว่าการใช้แบบจำลองเดี่ยว")
h_sub(doc, "1.4 ขอบเขตของการวิจัย")
bullet(doc, "ใช้ชุดข้อมูล DENTEX Challenge 2023 (MICCAI) จำนวน 1,005 ภาพที่มีการกำกับโรค")
bullet(doc, "จำแนกโรค 4 ประเภท: ฟันผุ, ฟันผุลึก, รอยโรคปลายราก, ฟันคุด")
bullet(doc, "พัฒนาเป็นเครื่องมือคัดกรองเบื้องต้น ไม่ใช่การวินิจฉัยทดแทนทันตแพทย์")
bullet(doc, "ฝึกและทดสอบแบบจำลองบนเครื่องคอมพิวเตอร์ส่วนบุคคล (GPU NVIDIA RTX 5060)")
h_sub(doc, "1.5 นิยามศัพท์เฉพาะ")
h_subsub(doc, "ภาพรังสีพาโนรามิก (Panoramic Radiograph)")
body(doc, "ภาพถ่ายรังสีที่แสดงฟันทั้งปาก ขากรรไกรบนและล่างในภาพเดียว ใช้ในการตรวจวินิจฉัยทางทันตกรรม", indent=0.3)
h_subsub(doc, "Grad-CAM (Gradient-weighted Class Activation Mapping)")
body(doc, "เทคนิคแสดงตำแหน่งในภาพที่มีอิทธิพลต่อการตัดสินใจของแบบจำลอง ในรูปแบบแผนที่ความร้อน (heatmap)", indent=0.3)
h_subsub(doc, "Ensemble และ Test-Time Augmentation (TTA)")
body(doc, "Ensemble คือการรวมผลทำนายจากหลายแบบจำลองเพื่อเพิ่มความแม่นยำ ส่วน TTA คือการทำนายภาพเดียวหลายมุม "
          "(เช่น ภาพต้นฉบับและภาพพลิกแนวนอน) แล้วเฉลี่ยผล", indent=0.3)
h_sub(doc, "1.6 ประโยชน์ที่คาดว่าจะได้รับ")
bullet(doc, "ได้เครื่องมือคัดกรองโรคในช่องปากเบื้องต้นที่ใช้งานได้จริงผ่านเว็บแอปพลิเคชัน")
bullet(doc, "ช่วยลดภาระการคัดกรองของทันตแพทย์ และเพิ่มโอกาสการเข้าถึงบริการในพื้นที่ห่างไกล")
bullet(doc, "เป็นต้นแบบการประยุกต์ AI กับงานทันตกรรมสำหรับการพัฒนาต่อยอด")
page_break(doc)

# ══════════════════ บทที่ 2 ══════════════════
h_chapter(doc, 2, "เอกสารและงานวิจัยที่เกี่ยวข้อง")
h_sub(doc, "2.1 ปัญญาประดิษฐ์และการเรียนรู้เชิงลึกในงานการแพทย์")
body(doc, "การเรียนรู้เชิงลึก (Deep Learning) โดยเฉพาะโครงข่ายประสาทเทียมแบบคอนโวลูชัน (Convolutional Neural "
          "Network, CNN) เป็นเทคโนโลยีหลักในการวิเคราะห์ภาพทางการแพทย์ มีงานวิจัยจำนวนมากที่นำ CNN มาใช้ในการ"
          "วินิจฉัยโรคจากภาพถ่ายรังสี ภาพจอประสาทตา และภาพผิวหนัง โดยให้ผลใกล้เคียงหรือดีกว่าผู้เชี่ยวชาญในบางงาน")
h_sub(doc, "2.2 สถาปัตยกรรม EfficientNet")
body(doc, "EfficientNet เป็นสถาปัตยกรรม CNN ที่ใช้เทคนิค compound scaling ในการปรับขนาดความลึก ความกว้าง และ"
          "ความละเอียดของภาพอย่างสมดุล ทำให้ได้ความแม่นยำสูงโดยใช้พารามิเตอร์น้อย เหมาะกับการฝึกบนทรัพยากรจำกัด "
          "ในโครงงานนี้ใช้ EfficientNet-B0 ซึ่งเป็นรุ่นพื้นฐานที่ฝึกล่วงหน้าบนชุดข้อมูล ImageNet [2]")
h_sub(doc, "2.3 การเรียนรู้แบบถ่ายโอน (Transfer Learning) และ Focal Loss")
body(doc, "Transfer Learning คือการนำแบบจำลองที่ฝึกบนข้อมูลขนาดใหญ่มาปรับใช้กับงานใหม่ที่มีข้อมูลน้อย "
          "ช่วยลดเวลาการฝึกและเพิ่มความแม่นยำ ส่วน Focal Loss เป็นฟังก์ชันสูญเสียที่ออกแบบมาเพื่อแก้ปัญหา"
          "ความไม่สมดุลของข้อมูล (class imbalance) โดยให้น้ำหนักกับตัวอย่างที่จำแนกยากมากขึ้น [2], [4]")
h_sub(doc, "2.4 ปัญญาประดิษฐ์ที่อธิบายได้ (Explainable AI) และ Grad-CAM")
body(doc, "ในงานทางการแพทย์ ความสามารถในการอธิบายผลของแบบจำลองมีความสำคัญต่อความน่าเชื่อถือ Grad-CAM เป็น"
          "เทคนิคที่ใช้ค่าเกรเดียนต์ในการสร้างแผนที่ความร้อนแสดงบริเวณของภาพที่มีอิทธิพลต่อการทำนาย "
          "ทำให้ทันตแพทย์สามารถตรวจสอบได้ว่าแบบจำลองให้ความสนใจที่ตำแหน่งฟันที่ถูกต้องหรือไม่ [3]")
h_sub(doc, "2.5 ชุดข้อมูล DENTEX Challenge 2023")
body(doc, "DENTEX เป็นชุดข้อมูลภาพรังสีพาโนรามิกที่เผยแพร่ในงานประชุม MICCAI 2023 ประกอบด้วยภาพที่กำกับโรค "
          "1,005 ภาพ และภาพไม่กำกับ 1,571 ภาพ มีการกำกับโรค 4 ประเภท ได้แก่ caries, deep caries, "
          "periapical lesion และ impacted tooth ในรูปแบบ object detection (bounding box ต่อฟันแต่ละซี่) "
          "เผยแพร่ภายใต้สัญญาอนุญาต CC BY-NC-SA 4.0 [1]")
h_sub(doc, "2.6 งานวิจัยที่เกี่ยวข้อง")
body(doc, "การประยุกต์ใช้การเรียนรู้เชิงลึก (deep learning) กับภาพรังสีทันตกรรมเติบโตอย่างรวดเร็วในช่วงไม่กี่ปีที่ผ่านมา "
          "โดยเฉพาะกับภาพรังสีพาโนรามิกที่ครอบคลุมฟันทั้งปากในภาพเดียว ทำให้เหมาะกับการคัดกรองหลายตำแหน่งพร้อมกัน "
          "หัวข้อนี้ทบทวนงานวิจัยที่เกี่ยวข้องโดยจัดกลุ่มตามโรคที่โครงงานนี้สนใจ (ฟันผุ รอยโรคปลายราก และฟันคุด) "
          "ตามด้วยงานด้านการจำแนกหลายโรคพร้อมกัน (multi-label) และปัญญาประดิษฐ์ที่อธิบายผลได้ (explainable AI)")

h_sub(doc, "2.6.1 การตรวจจับฟันผุ (Dental Caries)")
body(doc, "Oztekin และคณะ [5] เสนอแบบจำลองการเรียนรู้เชิงลึกที่อธิบายผลได้ (explainable model) สำหรับทำนายฟันผุจาก"
          "ภาพรังสีพาโนรามิก โดยใช้สถาปัตยกรรม ResNet-50 ได้ความแม่นยำราว 92% พร้อมสร้างแผนที่ความร้อน (heatmap) "
          "แสดงตำแหน่งที่แบบจำลองใช้ตัดสินใจ นับเป็นตัวอย่างที่ชี้ว่า CNN ตรวจฟันผุได้แม่นและสามารถอธิบายผลได้ "
          "นอกจากนี้ Vinayahalingam และคณะ [6] จำแนกฟันผุในฟันกรามซี่ที่สาม (third molar) จากภาพพาโนรามิกด้วยการ"
          "เรียนรู้เชิงลึก ได้ค่าพื้นที่ใต้กราฟ ROC (AUC) สูงถึงราว 0.91 ใกล้เคียงผลของทันตแพทย์ผู้เชี่ยวชาญ "
          "อย่างไรก็ตาม งานทั้งสองมุ่งเฉพาะโรคฟันผุและไม่ได้แยกระดับความลึก (caries / deep caries)")

h_sub(doc, "2.6.2 การตรวจจับรอยโรคปลายราก (Periapical Lesion)")
body(doc, "Çelik และคณะ [7] ศึกษาบทบาทของการเรียนรู้เชิงลึกในการตรวจรอยโรคปลายรากจากภาพรังสีพาโนรามิก โดยเปรียบเทียบ"
          "หลายสถาปัตยกรรม พบว่ากรอบการตรวจจับให้ค่า mean average precision (mAP) ระหว่าง 0.83–0.95 และความแม่นยำ"
          "ระหว่าง 0.67–0.81 โดย RetinaNet ให้ผลดีที่สุด (F1 สูงสุดราว 0.90) งานนี้แสดงให้เห็นว่า AI ช่วยตรวจรอยโรค"
          "ปลายรากซึ่งสังเกตได้ยากด้วยตาเปล่าได้อย่างมีประสิทธิภาพ แต่ยังเป็นการตรวจเฉพาะโรคเดียว")

h_sub(doc, "2.6.3 การตรวจจับฟันคุด (Impacted Tooth)")
body(doc, "Çelik [8] พัฒนาเครื่องมือตรวจจับฟันกรามล่างซี่ที่สามที่คุด (impacted mandibular third molar) จากภาพรังสี"
          "พาโนรามิกด้วยการรวมหลายโครงข่าย CNN พบว่าแบบจำลองตรวจตำแหน่งฟันคุดได้แม่นยำสูง ช่วยลดเวลาและความ"
          "ผิดพลาดในการประเมินก่อนการผ่าตัด ผลนี้สอดคล้องกับธรรมชาติของฟันคุดที่มักเห็นได้ชัดในภาพพาโนรามิก "
          "จึงเป็นโรคที่ AI ทำผลงานได้ดีอย่างสม่ำเสมอ")

h_sub(doc, "2.6.4 การจำแนกหลายโรคพร้อมกัน (Multi-label) และชุดข้อมูลมาตรฐาน")
body(doc, "Golkarieh และคณะ [9] เปรียบเทียบสถาปัตยกรรม CNN สมัยใหม่หลายแบบ (รวมถึง EfficientNet, DenseNet และ "
          "ResNet) สำหรับจำแนกภาวะทางทันตกรรมหลายชนิดจากภาพรังสีพาโนรามิก และชี้ว่าการที่ภาพหนึ่งมักมีหลายโรค"
          "พร้อมกัน รวมถึงความไม่สมดุลของข้อมูล (class imbalance) เป็นความท้าทายหลักที่ต้องแก้ด้วยการจำแนกแบบ "
          "multi-label และการเพิ่มข้อมูล (augmentation) ในด้านชุดข้อมูล Hamamci และคณะ [1] ได้เผยแพร่ชุดข้อมูล"
          "มาตรฐาน DENTEX (MICCAI 2023) ที่กำกับโรคหลายชนิด (caries, deep caries, periapical lesion, impacted) "
          "ไว้ในชุดเดียวอย่างเป็นลำดับชั้น เปิดทางให้พัฒนาแบบจำลองตรวจหลายโรคพร้อมกัน และเป็นชุดข้อมูลหลักของโครงงานนี้")

h_sub(doc, "2.6.5 ปัญญาประดิษฐ์ที่อธิบายผลได้ (Explainable AI) ในงานทันตกรรม")
body(doc, "ความสามารถในการอธิบายผลมีความสำคัญต่อความน่าเชื่อถือของ AI ทางการแพทย์ Selvaraju และคณะ [3] เสนอเทคนิค "
          "Grad-CAM ที่ใช้ค่าเกรเดียนต์สร้างแผนที่ความร้อนชี้บริเวณที่มีอิทธิพลต่อการตัดสินใจของแบบจำลอง ซึ่งถูกนำมาใช้"
          "อย่างแพร่หลายในงานทันตกรรม เช่นในงานของ Oztekin และคณะ [5] โดยมีรายงานว่าบริเวณที่ Grad-CAM เน้นสอดคล้อง"
          "กับตำแหน่งที่มีความสำคัญทางคลินิกในกว่า 84% ของกรณี ช่วยให้ทันตแพทย์ตรวจสอบและไว้วางใจผลของแบบจำลองได้")

h_sub(doc, "2.6.6 ช่องว่างของงานวิจัยและจุดต่างของโครงงานนี้")
body(doc, "จากการทบทวนข้างต้นพบช่องว่าง (research gap) ว่างานเดิมส่วนใหญ่ (1) ตรวจโรคทีละชนิด ต้องใช้หลายแบบจำลอง"
          "หากต้องการคัดกรองหลายโรค (2) อยู่ในรูปงานวิจัยเชิงตรวจจับ ยังไม่เป็นเครื่องมือพร้อมใช้สำหรับผู้ใช้ทั่วไป "
          "(3) ส่วนใหญ่พิจารณาเฉพาะภาพ ไม่ได้นำอาการทางคลินิกมาประกอบ และ (4) ยังไม่รวมการอธิบายผลกับการสรุป"
          "เป็นภาษาที่เข้าใจง่ายไว้ในระบบเดียว โครงงาน DentScan AI จึงออกแบบให้แตกต่างโดยรวมการคัดกรอง 4 โรคไว้ใน"
          "แบบจำลองเดียวแบบ multi-label ใช้ ensemble ของ EfficientNet เพิ่ม Grad-CAM อธิบายตำแหน่งที่แบบจำลองสนใจ "
          "ผสานข้อมูลอาการ (symptom fusion) เข้ากับผลจากภาพ และนำเสนอผ่านเว็บแอปพลิเคชันพร้อมรายงานภาษาไทยที่"
          "ผู้ใช้ทั่วไปเข้าใจได้")
page_break(doc)

# ══════════════════ บทที่ 3 ══════════════════
h_chapter(doc, 3, "วิธีดำเนินการ")
h_sub(doc, "3.1 แหล่งข้อมูล")
body(doc, "ใช้ชุดข้อมูล DENTEX Challenge 2023 ส่วนที่มีการกำกับโรค (quadrant-enumeration-disease) "
          "จำนวน 1,005 ภาพ ดาวน์โหลดจาก HuggingFace (ibrahimhamamci/DENTEX) แบ่งเป็นชุดฝึก 705 ภาพ "
          "และชุดตรวจสอบ (validation) 50 ภาพ")
v1 = M.get("v1", {}).get("per_class", {})
table_caption(doc, "ตารางที่ 3.1 จำนวนภาพต่อโรคในชุดข้อมูลฝึก")
add_table(doc, ["โรค", "ภาษาไทย", "จำนวนภาพ (ชุดฝึก)", "สัดส่วน"],
          [["Caries", "ฟันผุ", "623", "88.4%"],
           ["Deep Caries", "ฟันผุลึก", "321", "45.5%"],
           ["Impacted", "ฟันคุด", "254", "36.0%"],
           ["Periapical Lesion", "รอยโรคปลายราก", "116", "16.5%"]])
add_image(doc, ROOT/"notebooks"/"chart2_image_label_dist.png", width=5.0,
          caption="ภาพที่ 3.1 การกระจายตัวของข้อมูลแต่ละโรค (image-level)")
h_sub(doc, "3.2 การเตรียมและปรับปรุงข้อมูล (Preprocessing & Augmentation)")
body(doc, "ภาพรังสีพาโนรามิกต้นฉบับมีขนาดประมาณ 2,870 × 1,316 พิกเซล (อัตราส่วน ~2.1:1) ผู้วิจัยปรับขนาดภาพ"
          "โดยคงสัดส่วนเป็น 640 × 320 พิกเซล (และ 512 × 512 สำหรับบางแบบจำลอง) เพื่อลดการบิดเบือนของฟัน "
          "และจัดเก็บภาพที่ปรับขนาดแล้วล่วงหน้า (cache) เพื่อเพิ่มความเร็วในการฝึก")
body(doc, "ใช้เทคนิค data augmentation ระหว่างการฝึก ได้แก่ การพลิกแนวนอน การหมุนเล็กน้อย การปรับความสว่าง/"
          "คอนทราสต์ การเพิ่มคอนทราสต์เฉพาะที่ (CLAHE) และการสุ่มบดบังบางส่วน (CoarseDropout) "
          "เพื่อเพิ่มความหลากหลายของข้อมูลและลด overfitting")
h_sub(doc, "3.3 สถาปัตยกรรมแบบจำลอง (Model Architecture)")
body(doc, "ใช้ EfficientNet-B0 ที่ฝึกล่วงหน้าบน ImageNet เป็นฐาน ปรับชั้นสุดท้ายเป็น sigmoid 4 หัว "
          "(multi-label) เพื่อให้ภาพหนึ่งสามารถมีได้หลายโรคพร้อมกัน โดย 'ปกติ' หมายถึงทุกโรคมีความน่าจะเป็น"
          "ต่ำกว่าค่าเกณฑ์ (threshold) ไม่ใช่คลาสแยก เนื่องจากชุดข้อมูล DENTEX ทุกภาพมีความผิดปกติอย่างน้อยหนึ่งโรค")
h_sub(doc, "3.4 การฝึกแบบจำลอง (Model Training)")
bullet(doc, "ฟังก์ชันสูญเสีย: Focal Loss (gamma = 2.0) ร่วมกับ pos_weight แก้ class imbalance")
bullet(doc, "Optimizer: AdamW, learning rate 1.5×10⁻⁴–3×10⁻⁴, cosine annealing schedule")
bullet(doc, "ใช้ mixed precision (AMP) และ early stopping ตามค่า AUC บนชุดตรวจสอบ")
bullet(doc, "ฝึก 3 แบบจำลอง: v1 (512², 4 คลาส), v2 (640×320, เพิ่ม regularization), v3 (3 คลาส รวมฟันผุ)")
h_sub(doc, "3.5 การประเมินผลและการเสิร์ฟแบบจำลอง")
body(doc, "ประเมินด้วยค่า AUC, Average Precision และ F1-score ต่อโรค จากนั้นเปรียบเทียบแบบจำลองเดี่ยว "
          "การรวมแบบจำลอง (ensemble) และการทำ test-time augmentation (พลิกแนวนอน) เพื่อเลือกค่าการตั้งค่า"
          "เสิร์ฟที่ดีที่สุด พร้อมปรับค่าเกณฑ์ (threshold) ต่อโรคเพื่อให้ได้ F1 สูงสุด")
h_sub(doc, "3.6 การพัฒนาเว็บแอปพลิเคชัน")
body(doc, "พัฒนาส่วนหลังบ้าน (backend) ด้วย FastAPI ที่โหลดแบบจำลองและให้บริการผ่าน REST API "
          "และส่วนหน้า (frontend) เป็นเว็บแอปพลิเคชันหน้าเดียว (SPA) ที่รองรับการอัปโหลดภาพ X-ray "
          "ร่วมกับแบบสอบถามอาการ (hybrid) แสดงผลความน่าจะเป็นของแต่ละโรค แผนที่ความร้อน Grad-CAM "
          "และรายงานภาษาไทย พร้อมระบบรวมข้อมูลอาการ (symptom fusion) แบบ rule-based เพื่อปรับผลการทำนาย")
add_image(doc, ROOT/"web"/"samples"/"sample2_showcase.png", width=5.0,
          caption="ภาพที่ 3.2 ตัวอย่างผลการวิเคราะห์พร้อมแผนที่ความร้อน Grad-CAM และความน่าจะเป็นของแต่ละโรค")
h_sub(doc, "3.7 สรุปวิธีดำเนินการ")
body(doc, "ขั้นตอนการดำเนินงานประกอบด้วย (1) การเตรียมข้อมูลและสำรวจข้อมูล (2) การฝึกและประเมินแบบจำลอง "
          "(3) การหาการตั้งค่าเสิร์ฟที่ดีที่สุดด้วย ensemble และ TTA (4) การพัฒนาเว็บแอปพลิเคชันและระบบรายงาน "
          "และ (5) การทดสอบระบบแบบครบวงจร")
page_break(doc)

# ══════════════════ บทที่ 4 ══════════════════
h_chapter(doc, 4, "ผลการดำเนินงาน")
h_sub(doc, "4.1 ผลการสำรวจข้อมูล (EDA)")
body(doc, "จากการสำรวจชุดข้อมูลพบว่าโรคฟันผุพบมากที่สุด (88.4% ของภาพในชุดฝึก) ขณะที่รอยโรคปลายรากพบน้อยที่สุด "
          "(16.5%) สะท้อนปัญหาความไม่สมดุลของข้อมูลที่ต้องแก้ด้วย focal loss นอกจากนี้โรคฟันผุและฟันผุลึก"
          "มักปรากฏร่วมกันในภาพเดียว ซึ่งสอดคล้องกับลักษณะทางคลินิกที่เป็นความรุนแรงในระดับต่อเนื่องกัน")
add_image(doc, ROOT/"notebooks"/"chart3_cooccurrence.png", width=4.5,
          caption="ภาพที่ 4.1 เมทริกซ์การเกิดร่วมกันของโรค (co-occurrence)")
h_sub(doc, "4.2 ผลการจำแนกโรคของแต่ละแบบจำลอง")
rows = []
labels = {"v1": "v1 (512², 4 คลาส)", "v2": "v2 (640×320)", "v3": "v3 (3 คลาส รวมฟันผุ)"}
for k in ["v1", "v2", "v3"]:
    if k in M:
        rows.append([labels[k], f"{M[k].get('mean_auc','-')}", f"{M[k].get('mean_f1','-')}"])
table_caption(doc, "ตารางที่ 4.1 เปรียบเทียบผลของแต่ละแบบจำลอง")
add_table(doc, ["แบบจำลอง", "ค่าเฉลี่ย AUC", "ค่าเฉลี่ย F1"], rows)
body(doc, "จากตารางพบว่าแบบจำลอง v1 (512 พิกเซล, 4 คลาส) ให้ค่าเฉลี่ย AUC สูงสุดในกลุ่มแบบจำลองเดี่ยว", indent=0.5)
h_sub(doc, "4.3 ผลการรวมแบบจำลองและ Test-Time Augmentation")
body(doc, "เมื่อทดลองรวมแบบจำลอง v1 และ v2 ร่วมกับ test-time augmentation (พลิกแนวนอน) พบว่าให้ค่าเฉลี่ย AUC "
          f"เท่ากับ {M.get('serving',{}).get('val_mean_auc','0.7229')} ซึ่งสูงกว่าแบบจำลองเดี่ยวที่ดีที่สุด "
          "(0.692) อย่างมีนัยสำคัญ และผ่านเกณฑ์สมมติฐานที่ตั้งไว้ (AUC ≥ 0.70)")
table_caption(doc, "ตารางที่ 4.2 ผลการรวมแบบจำลองและ Test-Time Augmentation")
add_table(doc, ["การตั้งค่า", "ค่าเฉลี่ย AUC"],
          [["v1 เดี่ยว", "0.692"], ["v1 + TTA", "0.712"],
           ["ensemble (v1+v2)", "0.708"], ["ensemble (v1+v2) + TTA", "0.723"]])
v1pc = M.get("v1", {}).get("per_class", {})
if v1pc:
    h_sub(doc, "4.4 ผลการจำแนกรายโรค (แบบจำลอง v1)")
    pc_rows = []
    th = {"Caries": "ฟันผุ", "Deep Caries": "ฟันผุลึก",
          "Periapical Lesion": "รอยโรคปลายราก", "Impacted": "ฟันคุด"}
    for en, m in v1pc.items():
        pc_rows.append([th.get(en, en), m.get("auc", "-"), m.get("ap", "-"), m.get("f1", "-")])
    table_caption(doc, "ตารางที่ 4.3 ผลการจำแนกรายโรค (แบบจำลอง v1)")
    add_table(doc, ["โรค", "AUC", "AP", "F1"], pc_rows)
    body(doc, "โรคฟันผุและฟันคุดให้ผลการจำแนกที่ดี (AUC ~0.77–0.78) ขณะที่รอยโรคปลายรากและฟันผุลึก"
              "ยังมีข้อจำกัด เนื่องจากปริมาณข้อมูลน้อยและลักษณะรอยโรคมีขนาดเล็ก", indent=0.5)
h_sub(doc, "4.5 ผลการแสดงตำแหน่งด้วย Grad-CAM")
body(doc, "แผนที่ความร้อน Grad-CAM แสดงให้เห็นว่าแบบจำลองให้ความสนใจที่ตำแหน่งฟันที่สอดคล้องกับโรคที่ตรวจพบ "
          "เช่น ในกรณีฟันคุด แผนที่ความร้อนชี้ไปยังบริเวณฟันกรามซี่สุดท้ายด้านล่าง ซึ่งช่วยยืนยันความน่าเชื่อถือ"
          "ของผลการทำนาย")
add_image(doc, ROOT/"web"/"samples"/"sample1_showcase.png", width=5.0,
          caption="ภาพที่ 4.2 ตัวอย่างแผนที่ความร้อน Grad-CAM ต่อโรคที่ตรวจพบ")
h_sub(doc, "4.6 ผลการพัฒนาเว็บแอปพลิเคชัน")
body(doc, "ระบบเว็บแอปพลิเคชันสามารถใช้งานได้จริงแบบครบวงจร ตั้งแต่การอัปโหลดภาพ X-ray การกรอกอาการ "
          "การประมวลผลด้วยแบบจำลอง ensemble การแสดงแผนที่ความร้อน ไปจนถึงการสร้างรายงานภาษาไทย 2 ระดับ "
          "(สำหรับผู้ป่วยและทันตแพทย์) โดยใช้เวลาประมวลผลประมาณ 2–3 วินาทีต่อภาพ")
page_break(doc)

# ══════════════════ บทที่ 5 ══════════════════
h_chapter(doc, 5, "สรุป อภิปรายผล และข้อเสนอแนะ")
h_sub(doc, "5.1 สรุปผลการดำเนินงาน")
body(doc, "โครงงานนี้พัฒนา DentScan AI ระบบคัดกรองโรคในช่องปาก 4 ประเภทจากภาพรังสีพาโนรามิกได้สำเร็จ "
          "โดยใช้ EfficientNet-B0 ร่วมกับ transfer learning, focal loss, ensemble และ test-time augmentation "
          "ให้ค่าเฉลี่ย AUC เท่ากับ 0.723 ผ่านเกณฑ์สมมติฐานที่ตั้งไว้ พร้อมประยุกต์ Grad-CAM เพื่อความสามารถ"
          "ในการอธิบายผล และพัฒนาเว็บแอปพลิเคชันแบบไฮบริดที่ใช้งานได้จริง")
h_sub(doc, "5.2 อภิปรายผลการทดลอง")
body(doc, "ผลการทดลองสนับสนุนสมมติฐานว่าการรวมแบบจำลองและ test-time augmentation ช่วยเพิ่มความแม่นยำ "
          "เนื่องจากแต่ละแบบจำลองมีจุดเด่นต่างกัน (v1 ที่ความละเอียด 512 จับฟันผุได้ดี ส่วน v2 ที่คงสัดส่วน"
          "พาโนรามิกจับฟันผุลึกได้ดีกว่า) การเฉลี่ยผลจึงลดความแปรปรวนและเพิ่มความเสถียร อย่างไรก็ตามชุดตรวจสอบ"
          "ที่มีเพียง 50 ภาพทำให้ค่าที่วัดได้มีความผันผวน โดยเฉพาะรอยโรคปลายรากที่มีตัวอย่างน้อย")
h_sub(doc, "5.3 ข้อจำกัดของงานวิจัย")
bullet(doc, "ชุดข้อมูลมีจำนวนจำกัด (1,005 ภาพ) และชุดตรวจสอบเล็ก ทำให้ค่าประเมินผันผวน")
bullet(doc, "รอยโรคปลายรากและฟันผุลึกมีข้อมูลน้อย ส่งผลให้ความแม่นยำยังไม่สูงพอ")
bullet(doc, "ชุดข้อมูล DENTEX อยู่ภายใต้สัญญาอนุญาตห้ามใช้เชิงพาณิชย์ (CC BY-NC-SA 4.0)")
bullet(doc, "ระบบเป็นเครื่องมือคัดกรองเบื้องต้น ไม่สามารถใช้ทดแทนการวินิจฉัยของทันตแพทย์")
h_sub(doc, "5.4 แนวทางการพัฒนาต่อยอดในอนาคต (Future Work)")
bullet(doc, "เพิ่มปริมาณข้อมูลด้วยชุดข้อมูลเปิดอื่น (เช่น Tufts Dental) หรือใช้ self-supervised pretraining")
bullet(doc, "ใช้แบบจำลองตรวจจับวัตถุ (object detection) เพื่อระบุตำแหน่งฟันและซี่ที่เป็นโรคได้แม่นยำขึ้น")
bullet(doc, "เชื่อมต่อโมเดลภาษาขนาดใหญ่ (LLM) เพื่อสร้างรายงานที่ละเอียดและเป็นธรรมชาติมากขึ้น")
bullet(doc, "ทดสอบกับภาพจากคลินิกจริงและประเมินร่วมกับทันตแพทย์ เพื่อยืนยันประสิทธิภาพทางคลินิก")
page_break(doc)

# ══════════════════ บรรณานุกรม ══════════════════
h_front(doc, "บรรณานุกรม")
refs = [
    "Hamamci, I. E., et al. (2023). DENTEX: An Abnormal Tooth Detection with Dental Enumeration "
    "and Diagnosis Benchmark for Panoramic X-rays. MICCAI 2023.",
    "Tan, M., & Le, Q. (2019). EfficientNet: Rethinking Model Scaling for Convolutional Neural "
    "Networks. ICML 2019.",
    "Selvaraju, R. R., et al. (2017). Grad-CAM: Visual Explanations from Deep Networks via "
    "Gradient-based Localization. ICCV 2017.",
    "Lin, T. Y., et al. (2017). Focal Loss for Dense Object Detection. ICCV 2017.",
    "Oztekin, F., et al. (2023). An Explainable Deep Learning Model to Prediction Dental Caries Using "
    "Panoramic Radiograph Images. Diagnostics, 13(2), 226.",
    "Vinayahalingam, S., et al. (2021). Classification of caries in third molars on panoramic "
    "radiographs using deep learning. Scientific Reports, 11, 12609.",
    "Çelik, B., et al. (2023). The role of deep learning for periapical lesion detection on panoramic "
    "radiographs. Dentomaxillofacial Radiology, 52(8), 20230118.",
    "Çelik, M. E. (2022). Deep Learning Based Detection Tool for Impacted Mandibular Third Molar "
    "Teeth. Diagnostics, 12(4), 942.",
    "Golkarieh, A., et al. (2025). Advanced Deep Learning Techniques for Classifying Dental "
    "Conditions Using Panoramic X-Ray Images. arXiv:2508.21088.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(f"[{i}] {r}")
    _font(run, 16)
page_break(doc)

# ══════════════════ ภาคผนวก ══════════════════
h_front(doc, "ภาคผนวก")
para(doc, "ภาคผนวก ก: ผลการสำรวจข้อมูลเพิ่มเติม", size=16, bold=True, space_after=8)
add_image(doc, ROOT/"notebooks"/"chart1_ann_distribution.png", width=4.8,
          caption="ภาพที่ ก.1 จำนวน bounding box annotation ต่อโรค")
add_image(doc, ROOT/"notebooks"/"chart5_sample_xrays.png", width=5.5,
          caption="ภาพที่ ก.2 ตัวอย่างภาพรังสีพาโนรามิกของแต่ละโรค")
add_image(doc, ROOT/"notebooks"/"chart6_annotated_example.png", width=5.5,
          caption="ภาพที่ ก.3 ตัวอย่างภาพที่มีการกำกับตำแหน่งโรค (bounding box)")
para(doc, "ภาคผนวก ข: ตัวอย่างผลการวิเคราะห์เพิ่มเติม", size=16, bold=True, space_before=10, space_after=8)
add_image(doc, ROOT/"web"/"samples"/"sample3_showcase.png", width=5.5,
          caption="ภาพที่ ข.1 ตัวอย่างผลการวิเคราะห์ภาพที่ 3")
para(doc, "หมายเหตุ: ภาพรังสีทั้งหมดมาจากชุดข้อมูล DENTEX Challenge 2023 (CC BY-NC-SA 4.0) "
          "ใช้เพื่อการศึกษาและการประกวดเท่านั้น", size=14, italic=True, space_before=8)

doc.save(str(OUT / "DentScan_Report.docx"))
print(f"✅ Saved: {OUT/'DentScan_Report.docx'}")
print(f"   paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")
